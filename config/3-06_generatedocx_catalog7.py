import openpyxl
from docx import Document
import os

# 定义Excel文件和Word模板路径
excel_file_path = r"./template/minor/export.xlsx"
word_template_path = r"./template/minor/catalog15/temp.docx"

# 打开Excel文件
wb = openpyxl.load_workbook(excel_file_path)
sheet = wb.active

# 获取A2单元格的内容作为文件名
output_file_name = sheet['A92'].value
if output_file_name is not None:  # 添加条件语句检查A2是否为空
    # 构建完整的输出文件路径
    output_word_path = os.path.join(os.path.dirname(word_template_path), f"{output_file_name}.docx")

    # 打开Word文档
    doc = Document(word_template_path)

    # 获取标题（第一行的标题）
    headers = [cell.value for cell in sheet[1]]

    # 替换Word文档中的内容
    for i in range(92, 107):
        for col, header in enumerate(headers):
            cell_value = sheet.cell(row=i, column=col + 1).value

            # 构造要查找和替换的字符串
            find_str = f"【{header}{i-91}】"  # 注意索引要减1

            # 如果单元格的值为None，则将replace_str设置为空白字符串
            replace_str = str(cell_value) if cell_value is not None else ""

            # 遍历Word文档中的段落
            for paragraph in doc.paragraphs:
                if find_str in paragraph.text:
                    # 替换文本
                    paragraph.text = paragraph.text.replace(find_str, replace_str)

            # 遍历Word文档中的表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find_str in cell.text:
                            # 替换表格中的文本
                            cell.text = cell.text.replace(find_str, replace_str)

    # 保存修改后的Word文档
    doc.save(output_word_path)

    print(f"文件已保存到: {output_file_name}")
else:
    print("A92单元格的内容为空")

# 关闭Excel文件
wb.close()
