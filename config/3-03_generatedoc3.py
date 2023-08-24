import os
from openpyxl import load_workbook
from docx import Document

def combine_runs(paragraph):
    combined_text = ""
    for run in paragraph.runs:
        combined_text += run.text
    return combined_text

def replace_in_template(template_path, data_row):
    doc = Document(template_path)
    for key, value in data_row.items():
        placeholder = f"【{key}】"
        for paragraph in doc.paragraphs:
            if placeholder in combine_runs(paragraph):
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in combine_runs(paragraph):
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    return doc

def main():
    excel_file_path = r".\template\main\export.xlsx"
    word_template_path = r".\template\main\cover\temp.docx"

    wb = load_workbook(excel_file_path)
    ws = wb.active

    # Assuming the first row contains the header/title names
    headers = [cell.value for cell in ws[1]]

    # Skip the first row (header) and process the remaining rows
    for row in ws.iter_rows(min_row=2, values_only=True):
        data_row = dict(zip(headers, row))
        file_name = f"{row[0]}.docx"  # Use the value from A column as the filename
        new_doc = replace_in_template(word_template_path, data_row)

        # Save the new Word document in the same directory as the template
        output_path = os.path.join(os.path.dirname(word_template_path), file_name)
        new_doc.save(output_path)

if __name__ == "__main__":
    main()
